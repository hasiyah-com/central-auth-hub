#!/usr/bin/env python3
"""Build Thai DOCX reports with complete metric charts for RBA V4-V6."""

from __future__ import annotations

import argparse
import json
import math
import os
from pathlib import Path
from typing import Any

os.environ.setdefault("MPLCONFIGDIR", "/tmp/rba-matplotlib")

import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import precision_recall_curve, roc_curve


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
VERSIONS = {
    "v4": {
        "title": "RBA Adversarial Benchmark V4",
        "subtitle": "การโจมตีหลายระยะที่เลียนแบบพฤติกรรมผู้ใช้",
        "root": ROOT / "results" / "adversarial_v4",
        "stage": "sequence_v4",
        "gate_key": "ready_for_adversarial_shadow",
        "change": "เพิ่ม attack campaign 6 กลุ่ม กลุ่มละ 4 phases และ prototype evidence ledger",
    },
    "v5": {
        "title": "RBA Sequence Model V5",
        "subtitle": "One-class sequence model ที่ฝึกจาก trusted normal เท่านั้น",
        "root": ROOT / "results" / "sequence_model_v5",
        "stage": "sequence_model_v5",
        "gate_key": "ready_for_system_integration_shadow",
        "change": "เพิ่ม rolling sequence features และ normal-only Isolation Forest calibration",
    },
    "v6": {
        "title": "RBA Supervised Sequence Model V6",
        "subtitle": "Supervised hybrid พร้อม train/calibration/test attack seed แยกกัน",
        "root": ROOT / "results" / "supervised_sequence_v6",
        "stage": "supervised_sequence_v6",
        "gate_key": "ready_for_system_integration_shadow",
        "change": "เพิ่ม Random Forest sequence layer และ independent normal calibration",
    },
    "v7": {
        "title": "RBA Deployable Shadow Bundle V7",
        "subtitle": "Serialization parity, runtime contract และ latency validation",
        "root": ROOT / "results" / "deployable_bundle_v7",
        "stage": "supervised_sequence_v6",
        "gate_key": "ready_for_system_shadow_load",
        "change": "บรรจุ V6 เป็น joblib bundle พร้อม runtime ที่ตรวจ schema และไม่สามารถ enforce decision",
    },
}

COLORS = ["#2E74B5", "#4C9F70", "#D99A2B", "#A44A3F", "#6B5B95", "#607D8B"]
THAI_FONT_FILE = Path(__file__).resolve().parents[2] / "report_work" / "font_home" / "fonts" / "NotoSansThai.ttf"
FONT = "Noto Sans Thai"
if THAI_FONT_FILE.exists():
    font_manager.fontManager.addfont(str(THAI_FONT_FILE))
plt.rcParams["font.family"] = FONT


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_document(doc: Document, version: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Title", 24, "0B2545", 0, 6),
        ("Subtitle", 13, "4B5563", 0, 14),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.text = f"RBA USER LEARNING CURVE | {version.upper()}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer.paragraphs[0]
    footer.text = "Synthetic isolated experiment - Shadow evaluation only"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)


def _save_plot(path: Path) -> None:
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()


def _bar(path: Path, title: str, labels: list[str], values: list[float], percent=True, horizontal=False) -> None:
    plt.figure(figsize=(8.2, 4.2))
    positions = np.arange(len(labels))
    display = np.asarray(values) * (100.0 if percent else 1.0)
    if horizontal:
        bars = plt.barh(positions, display, color=COLORS[: len(labels)])
        plt.yticks(positions, labels)
        plt.gca().invert_yaxis()
        for bar, value in zip(bars, display):
            plt.text(value + max(display.max(), 1) * 0.01, bar.get_y() + bar.get_height() / 2, f"{value:.2f}{'%' if percent else ''}", va="center", fontsize=8)
    else:
        bars = plt.bar(positions, display, color=COLORS[: len(labels)])
        plt.xticks(positions, labels, rotation=20, ha="right")
        for bar, value in zip(bars, display):
            plt.text(bar.get_x() + bar.get_width() / 2, value + max(display.max(), 1) * 0.015, f"{value:.2f}{'%' if percent else ''}", ha="center", fontsize=8)
    plt.title(title, loc="left", fontweight="bold")
    plt.ylabel("Percent" if percent else "Value")
    plt.grid(axis="x" if horizontal else "y", alpha=0.2)
    _save_plot(path)


def build_charts(version: str, config: dict[str, Any], out: Path) -> dict[str, Path]:
    root = config["root"]
    stage_df = pd.read_csv(root / "stage_run_results.csv")
    attack_df = pd.read_csv(root / ("attack_sequence_run_results.csv" if version != "v4" else "attack_sequence_run_results.csv"))
    selected = stage_df[stage_df.stage.eq(config["stage"])]
    attack_selected = attack_df[attack_df.stage.eq(config["stage"])]
    out.mkdir(parents=True, exist_ok=True)
    charts: dict[str, Path] = {}

    performance = [name for name in ("precision", "event_challenge_recall", "f1", "roc_auc", "pr_auc") if name in selected]
    if performance:
        charts["performance"] = out / "01_performance.png"
        _bar(charts["performance"], "Overall classification metrics", performance, [float(selected[name].mean()) for name in performance])

    fprs = [name for name in ("challenge_fpr", "warn_fpr", "block_fpr") if name in selected]
    charts["fpr"] = out / "02_false_positive.png"
    _bar(charts["fpr"], "Normal false-positive rates", fprs, [float(selected[name].mean()) for name in fprs])

    sequence_metrics = [name for name in ("sequence_detection_rate", "preobjective_detection_rate", "objective_detection_rate") if name in selected]
    charts["sequence"] = out / "03_sequence_metrics.png"
    _bar(charts["sequence"], "Sequence-level detection", sequence_metrics, [float(selected[name].mean()) for name in sequence_metrics])

    family = attack_selected.groupby("attack_type").sequence_detection_rate.mean().sort_values()
    charts["attack"] = out / "04_attack_detection.png"
    _bar(charts["attack"], "Sequence detection by attack family", list(family.index), list(family.values), horizontal=True)

    size = selected.groupby("dataset_size").mean(numeric_only=True).reset_index()
    charts["learning"] = out / "05_learning_curve.png"
    plt.figure(figsize=(8.2, 4.2))
    for metric, color in zip(("sequence_detection_rate", "preobjective_detection_rate", "event_challenge_recall"), COLORS):
        if metric in size:
            plt.plot(size.dataset_size, size[metric] * 100.0, marker="o", label=metric, color=color)
    plt.xscale("log")
    plt.xlabel("Normal events per user (log scale)")
    plt.ylabel("Percent")
    plt.title("Learning curve by dataset size", loc="left", fontweight="bold")
    plt.legend(fontsize=8)
    plt.grid(alpha=0.2)
    _save_plot(charts["learning"])

    charts["fpr_size"] = out / "06_fpr_by_size.png"
    plt.figure(figsize=(8.2, 4.2))
    for metric, color in zip(("challenge_fpr", "warn_fpr"), (COLORS[3], COLORS[2])):
        if metric in size:
            plt.plot(size.dataset_size, size[metric] * 100.0, marker="o", label=metric, color=color)
    plt.xscale("log")
    plt.xlabel("Normal events per user (log scale)")
    plt.ylabel("Percent")
    plt.title("Normal friction by dataset size", loc="left", fontweight="bold")
    plt.legend(fontsize=8)
    plt.grid(alpha=0.2)
    _save_plot(charts["fpr_size"])

    scenario = selected.groupby("normal_scenario").mean(numeric_only=True)
    charts["nat"] = out / "07_nat_scenarios.png"
    labels = list(scenario.index)
    x = np.arange(len(labels))
    width = 0.35
    plt.figure(figsize=(8.2, 4.2))
    plt.bar(x - width / 2, scenario.sequence_detection_rate * 100, width, label="Sequence detection", color=COLORS[0])
    plt.bar(x + width / 2, scenario.challenge_fpr * 100, width, label="Challenge FPR", color=COLORS[3])
    plt.xticks(x, labels)
    plt.ylabel("Percent")
    plt.title("Normal staggered vs NAT burst", loc="left", fontweight="bold")
    plt.legend(fontsize=8)
    plt.grid(axis="y", alpha=0.2)
    _save_plot(charts["nat"])

    ttd = attack_selected.groupby("attack_type").median_time_to_detect_phase.mean().sort_values()
    charts["ttd"] = out / "08_time_to_detect.png"
    _bar(charts["ttd"], "Mean median phase-to-detect by attack", list(ttd.index), list(ttd.values), percent=False, horizontal=True)

    if "mean_attack_mimicry_distance" in attack_selected:
        mimic = attack_selected.groupby("attack_type").mean_attack_mimicry_distance.mean().sort_values()
        charts["mimicry"] = out / "09_mimicry_distance.png"
        _bar(charts["mimicry"], "Mean profile-distance by attack", list(mimic.index), list(mimic.values), percent=False, horizontal=True)

    predictions_path = root / "predictions.csv"
    if predictions_path.exists():
        predictions = pd.read_csv(predictions_path)
        predictions = predictions[predictions.stage.eq(config["stage"])]
        y = predictions.label.astype(int).to_numpy()
        scores = predictions.anomaly_score.astype(float).to_numpy()
        fpr, tpr, _ = roc_curve(y, scores)
        precision, recall, _ = precision_recall_curve(y, scores)
        charts["roc"] = out / "10_roc_curve.png"
        plt.figure(figsize=(7.2, 4.4)); plt.plot(fpr, tpr, color=COLORS[0], linewidth=2); plt.plot([0, 1], [0, 1], "--", color="#94A3B8"); plt.xlabel("False positive rate"); plt.ylabel("True positive rate"); plt.title("ROC curve - all held-out predictions", loc="left", fontweight="bold"); plt.grid(alpha=0.2); _save_plot(charts["roc"])
        charts["pr"] = out / "11_pr_curve.png"
        plt.figure(figsize=(7.2, 4.4)); plt.plot(recall, precision, color=COLORS[1], linewidth=2); plt.xlabel("Recall"); plt.ylabel("Precision"); plt.title("Precision-Recall curve - all held-out predictions", loc="left", fontweight="bold"); plt.grid(alpha=0.2); _save_plot(charts["pr"])
        decision_positive = predictions.decision.isin(["challenge", "block"]).astype(int)
        tp = int(((y == 1) & (decision_positive == 1)).sum()); fp = int(((y == 0) & (decision_positive == 1)).sum()); tn = int(((y == 0) & (decision_positive == 0)).sum()); fn = int(((y == 1) & (decision_positive == 0)).sum())
        charts["confusion"] = out / "12_confusion_matrix.png"
        matrix = np.asarray([[tn, fp], [fn, tp]])
        plt.figure(figsize=(5.4, 4.6)); plt.imshow(matrix, cmap="Blues"); plt.xticks([0, 1], ["Pred normal", "Pred attack"]); plt.yticks([0, 1], ["True normal", "True attack"]); plt.title("Confusion matrix - challenge+", loc="left", fontweight="bold")
        for i in range(2):
            for j in range(2): plt.text(j, i, f"{matrix[i,j]:,}", ha="center", va="center", fontsize=12, color="white" if matrix[i,j] > matrix.max()/2 else "black")
        _save_plot(charts["confusion"])
        charts["distribution"] = out / "13_score_distribution.png"
        plt.figure(figsize=(8.2, 4.2)); plt.hist(scores[y == 0], bins=30, alpha=0.65, label="Normal", color=COLORS[0], density=True); plt.hist(scores[y == 1], bins=30, alpha=0.55, label="Attack", color=COLORS[3], density=True); plt.xlabel("Anomaly score"); plt.ylabel("Density"); plt.title("Score distribution", loc="left", fontweight="bold"); plt.legend(); plt.grid(alpha=0.15); _save_plot(charts["distribution"])
    latency_path = root / "latency_results.csv"
    if latency_path.exists():
        latency = pd.read_csv(latency_path).iloc[0]
        charts["latency"] = out / "14_latency.png"
        _bar(
            charts["latency"],
            "Single-login shadow inference latency",
            ["p50", "p95", "p99", "max"],
            [float(latency.p50_ms), float(latency.p95_ms), float(latency.p99_ms), float(latency.max_ms)],
            percent=False,
        )
    return charts


def _add_picture(doc: Document, path: Path, caption: str, interpretation: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(6.55))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)
    note = doc.add_paragraph(interpretation)
    note.paragraph_format.space_after = Pt(10)


def _metric_value(selected: pd.DataFrame, name: str) -> str:
    if name not in selected:
        return "N/A"
    value = float(selected[name].mean())
    return f"{value * 100:.2f}%"


def build_report(version: str, config: dict[str, Any], charts: dict[str, Path], output: Path) -> None:
    root = config["root"]
    stage_df = pd.read_csv(root / "stage_run_results.csv")
    attack_df = pd.read_csv(root / "attack_sequence_run_results.csv")
    selected = stage_df[stage_df.stage.eq(config["stage"])]
    attack_selected = attack_df[attack_df.stage.eq(config["stage"])]
    gate = json.loads((root / "release_gate.json").read_text(encoding="utf-8"))
    ready = bool(gate.get(config["gate_key"], False))

    doc = Document()
    _style_document(doc, version)
    p = doc.add_paragraph(style="Title")
    p.add_run(config["title"])
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(config["subtitle"])
    meta = doc.add_paragraph()
    meta.add_run("ขอบเขต: ").bold = True
    meta.add_run("12 alias profiles | IP 192.168.10.1 | ไม่มี Geo | 6 sizes | 5 seeds | 2 NAT scenarios")

    callout = doc.add_table(rows=1, cols=1)
    _set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    _set_cell_fill(cell, "E8F5E9" if ready else "FDECEC")
    status = "ผ่านเกณฑ์สำหรับ isolated shadow integration" if ready else "ยังไม่ผ่านเกณฑ์สำหรับ shadow integration"
    run = cell.paragraphs[0].add_run(f"คำตัดสิน: {status}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F6F43" if ready else "9B1C1C")
    cell.add_paragraph("Enforcement ยังคงเป็น false เสมอ จนกว่าจะผ่าน production replay, latency/serialization, monitoring, rollback และ canary")

    doc.add_heading("1. สรุปผู้บริหาร", level=1)
    doc.add_paragraph(config["change"] + " ผลลัพธ์ในรายงานนี้มาจาก synthetic isolated replay และไม่แตะ production database หรือ production model")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Metric", "Observed", "ความหมาย"]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        _set_cell_fill(table.cell(0, index), "F2F4F7")
    rows = [
        ("Precision", _metric_value(selected, "precision"), "สัดส่วน challenge+ ที่เป็น attack จริง"),
        ("Event recall", _metric_value(selected, "event_challenge_recall"), "สัดส่วน attack phases ที่ถูก challenge"),
        ("F1", _metric_value(selected, "f1"), "สมดุล Precision และ Recall"),
        ("ROC-AUC", _metric_value(selected, "roc_auc"), "ความสามารถจัดอันดับ normal/attack"),
        ("PR-AUC", _metric_value(selected, "pr_auc"), "คุณภาพการจัดอันดับเมื่อคลาสไม่สมดุล"),
        ("Sequence detection", _metric_value(selected, "sequence_detection_rate"), "campaign ที่ถูกจับอย่างน้อยหนึ่ง phase"),
        ("Pre-objective detection", _metric_value(selected, "preobjective_detection_rate"), "campaign ที่ถูกจับก่อน objective"),
        ("Challenge FPR", _metric_value(selected, "challenge_fpr"), "normal ที่ถูก challenge โดยไม่จำเป็น"),
        ("Warn FPR", _metric_value(selected, "warn_fpr"), "normal ที่ถูก warn หรือสูงกว่า"),
    ]
    for metric, value, meaning in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = metric, value, meaning
    _set_table_geometry(table, [2200, 1800, 5360])

    doc.add_heading("2. Data contract และวิธีทดลอง", level=1)
    bullets = [
        "ข้อมูลปกติแบ่งตามเวลา 80:20 และรักษา trusted-history allowlist = allow, mfa_passed เท่านั้น",
        "Admin ต้อง MFA ทุกครั้ง และ decision ที่เป็น warn/challenge/block รวม shadow labels ไม่อัปเดต trusted profile",
        "Attack เป็น 4 phases ใช้ device/UA/OS ที่เคยเห็น อยู่ในสิทธิ์ subsystem และต่ำกว่า Rule action floors",
        "Normal staggered และ Normal NAT burst ใช้ IP เดียวกัน 192.168.10.1 โดยไม่มี Geo",
        "ทุก release gate อนุญาตสูงสุดเพียง isolated shadow integration; ไม่อนุญาต enforcement",
    ]
    for text in bullets:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3. ผลการวัดประสิทธิภาพ", level=1)
    chart_notes = {
        "performance": ("กราฟ 1: Classification metrics", "ใช้ดูคุณภาพการจัดอันดับและสมดุลระหว่างการจับ attack กับความแม่นของ alert"),
        "fpr": ("กราฟ 2: False-positive rates", "Challenge FPR เป็น gate หลักด้าน friction; Warn FPR แสดงภาระการตรวจสอบเพิ่มเติม"),
        "sequence": ("กราฟ 3: Sequence detection", "Pre-objective detection สำคัญกว่าการจับได้หลัง objective เพราะสะท้อนความสามารถป้องกันล่วงหน้า"),
        "attack": ("กราฟ 4: ผลแยกตาม attack family", "ค่าต่ำสุดของทุก family ถูกใช้เป็น gate เพื่อไม่ให้ค่าเฉลี่ยซ่อนจุดอ่อน"),
        "learning": ("กราฟ 5: Learning curve", "เปรียบเทียบว่าประวัติ 10 ถึง 5,000 เหตุการณ์ต่อผู้ใช้ทำให้ประสิทธิภาพเปลี่ยนอย่างไร"),
        "fpr_size": ("กราฟ 6: Friction ตามขนาดข้อมูล", "ใช้ตรวจว่าประวัติที่ยาวขึ้นทำให้คะแนนสะสมจนเกิด false positive หรือไม่"),
        "nat": ("กราฟ 7: Staggered เทียบ NAT burst", "ใช้ยืนยันว่า shared private IP และการ login พร้อมกันไม่ทำให้ระบบลำเอียงเกินเกณฑ์"),
        "ttd": ("กราฟ 8: Time to detect", "ค่า 1 หมายถึงตรวจพบตั้งแต่ phase แรก; ค่า 5 หมายถึงไม่พบภายใน 4 phases"),
        "mimicry": ("กราฟ 9: Profile-distance", "ค่าต่ำแปลว่า attack ใกล้กับพฤติกรรมปกติและตรวจจับยากกว่า"),
        "roc": ("กราฟ 10: ROC curve", "วัดการแยกอันดับทุก threshold แต่ต้องอ่านร่วมกับ FPR ที่ threshold ใช้งานจริง"),
        "pr": ("กราฟ 11: Precision-Recall curve", "เหมาะกับข้อมูล attack/normal ที่ไม่สมดุลและสะท้อนคุณภาพ alert"),
        "confusion": ("กราฟ 12: Confusion matrix", "สรุป TP, FP, TN และ FN ที่ decision challenge+"),
        "distribution": ("กราฟ 13: Score distribution", "ตรวจระดับการซ้อนทับของคะแนน normal และ attack รวมถึงความเสี่ยงจาก threshold drift"),
        "latency": ("กราฟ 14: Runtime latency", "วัด single-login inference หลัง warm-up; gate กำหนด p95 ไม่เกิน 20 ms และ p99 ไม่เกิน 35 ms"),
    }
    for key in ("performance", "fpr", "sequence", "attack", "learning", "fpr_size", "nat", "ttd", "mimicry", "roc", "pr", "confusion", "distribution", "latency"):
        if key in charts:
            caption, note = chart_notes[key]
            _add_picture(doc, charts[key], caption, note)

    doc.add_heading("4. ผลแยกตาม Attack family", level=1)
    family = attack_selected.groupby("attack_type").mean(numeric_only=True).reset_index()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, value in enumerate(("Attack", "Sequence", "Pre-objective", "Median phase")):
        table.cell(0, index).text = value
        _set_cell_fill(table.cell(0, index), "F2F4F7")
    for _, row in family.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row.attack_type)
        cells[1].text = f"{row.sequence_detection_rate * 100:.2f}%"
        cells[2].text = f"{row.preobjective_detection_rate * 100:.2f}%"
        cells[3].text = f"{row.median_time_to_detect_phase:.2f}"
    _set_table_geometry(table, [3960, 1800, 1800, 1800])

    doc.add_heading("5. Release gate", level=1)
    gate_table = doc.add_table(rows=1, cols=3)
    gate_table.style = "Table Grid"
    for index, value in enumerate(("Check", "ผล", "สถานะ")):
        gate_table.cell(0, index).text = value
        _set_cell_fill(gate_table.cell(0, index), "F2F4F7")
    for name, passed in gate.get("checks", {}).items():
        cells = gate_table.add_row().cells
        cells[0].text = name
        observed = gate.get("observed", {}).get(name.replace("_ge_0_90", "").replace("_le_0_003", ""), "")
        cells[1].text = f"{observed}" if observed != "" else "ดู observed summary"
        cells[2].text = "PASS" if passed else "FAIL"
        _set_cell_fill(cells[2], "E8F5E9" if passed else "FDECEC")
    _set_table_geometry(gate_table, [5000, 2500, 1860])

    doc.add_heading("6. ข้อสรุปและขั้นถัดไป", level=1)
    if ready:
        doc.add_paragraph("เวอร์ชันนี้ผ่าน synthetic shadow gate แต่ยังไม่ใช่ production-ready enforcement การนำเข้าใช้กับระบบต้องเริ่มจาก shadow adapter ที่ไม่เปลี่ยน decision จริง เก็บ latency, drift, decision disagreement และ rollback telemetry")
    else:
        doc.add_paragraph("เวอร์ชันนี้ยังไม่ควรเชื่อมกับ request path แม้ใน shadow mode จนกว่าจะลด false positive หรือเพิ่มการตรวจจับ attack family ที่ต่ำกว่า gate โดยห้ามลด threshold เพื่อบังคับให้ผ่านโดยไม่มี independent calibration")
    for text in (
        "Replay anonymized production logs แบบ point-in-time และห้าม export identity จริง",
        "ตรวจ serialization, feature parity, p95/p99 latency, memory และ timeout behavior",
        "ติดตั้ง drift monitoring, threshold versioning, audit reasons และ one-command rollback",
        "เริ่ม canary เฉพาะหลัง shadow metrics ผ่านต่อเนื่องและมี security review",
    ):
        doc.add_paragraph(text, style="List Number")

    doc.core_properties.title = config["title"]
    doc.core_properties.subject = "RBA model experiment report"
    doc.core_properties.author = "RBA Experiment Automation"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--versions", nargs="+", choices=VERSIONS, default=list(VERSIONS))
    parser.add_argument("--output", type=Path, default=REPORT_DIR)
    args = parser.parse_args()
    for version in args.versions:
        config = VERSIONS[version]
        chart_dir = args.output / "charts" / version
        charts = build_charts(version, config, chart_dir)
        output = args.output / f"RBA_Model_Experiment_{version.upper()}_TH.docx"
        build_report(version, config, charts, output)
        print(output)


if __name__ == "__main__":
    main()
